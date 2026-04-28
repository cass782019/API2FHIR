"""Gerador de documento DOCX — FHIR-Forge Arquitetura.

Uso:
    uv run python scripts/generate_docx.py
Saída:
    docs/fhir-forge-arquitetura.docx  (~35 páginas, PT-BR)
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[attr-defined]
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Paleta ───────────────────────────────────────────────────────────────────
C_NAVY  = RGBColor(0x0D, 0x21, 0x37)
C_BLUE  = RGBColor(0x1E, 0x6F, 0xDB)
C_GREEN = RGBColor(0x13, 0xA1, 0x63)
C_GREY  = RGBColor(0x6B, 0x72, 0x80)
C_AMBER = RGBColor(0xD9, 0x77, 0x06)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_LBLUE = RGBColor(0xDB, 0xEA, 0xFB)
C_LGREY = RGBColor(0xF0, 0xF4, 0xFA)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = C_BLUE

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = C_NAVY

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = C_GREY


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, *,
       bold: bool = False, italic: bool = False,
       color: RGBColor | None = None, size: int = 11,
       align: object = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    para = doc.add_paragraph()
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def _bullet(doc: Document, items: list[str], *, size: int = 11) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        run.font.size = Pt(size)
        run.font.name = "Calibri"


def _set_cell_bg(cell: object, hex_str: str) -> None:
    """Define cor de fundo de célula via XML."""
    tc = cell._tc
    tcp = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str)
    tcp.append(shd)


def _table(doc: Document, headers: list[str], rows: list[list[str]],
           col_widths: list[float] | None = None,
           header_bg: str = "0D2137", header_fg: RGBColor = C_WHITE,
           alt_bg: str = "F0F4FA") -> None:
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"

    # Cabeçalho
    hrow = tbl.rows[0]
    for j, hdr in enumerate(headers):
        cell = hrow.cells[j]
        _set_cell_bg(cell, header_bg)
        para = cell.paragraphs[0]
        run = para.add_run(hdr)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = header_fg

    # Dados
    for i, row in enumerate(rows):
        trow = tbl.rows[i + 1]
        bg = alt_bg if i % 2 == 0 else "FFFFFF"
        for j, cell_txt in enumerate(row):
            cell = trow.cells[j]
            _set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_txt))
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    if col_widths:
        for i, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])

    doc.add_paragraph()


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def _hr(doc: Document) -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E6FDB")
    pb.append(bottom)
    pPr.append(pb)


# ── Conteúdo ─────────────────────────────────────────────────────────────────

def _capa(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    _p(doc, "FHIR-Forge", bold=True, size=36, color=C_BLUE,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "Documento de Arquitetura", bold=True, size=22, color=C_NAVY,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _p(doc,
       "Pipeline Híbrido LLM + SUSHI para Interoperabilidade em Saúde Digital\n"
       "Conversão Automática de Especificações de API em Recursos FHIR R4",
       italic=True, size=14, color=C_GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    meta = [
        ("Versão",  "0.1.0"),
        ("Data",    "Abril 2026"),
        ("Licença", "Apache-2.0"),
        ("Autor",   "Cassiano Moralles"),
        ("Python",  "≥ 3.12"),
        ("FHIR",    "R4 + BR Core 1.0.0"),
        ("Status",  "9/9 fases completas — 289 unit + 3 integration + 7 regression + 6 e2e, 91% cobertura — tag v2.1"),
    ]
    _table(doc, ["Campo", "Valor"], meta, col_widths=[2.0, 4.5])
    _hr(doc)
    _h(doc, "Resumo Executivo", level=2)
    _p(doc,
       "O FHIR-Forge é um sistema de software que automatiza a conversão de especificações "
       "de APIs REST (OpenAPI/Swagger) em recursos FHIR R4 conformes ao padrão BR Core e à "
       "plataforma federal RNDS (Rede Nacional de Dados em Saúde). Utiliza um pipeline híbrido "
       "composto por um agente LangGraph, modelos de linguagem de grande escala (LLMs) e um "
       "servidor FHIR (HAPI FHIR), com suporte completo a privacidade por design (LGPD).")
    _p(doc,
       "O sistema endereça o desafio de interoperabilidade em saúde no Brasil: centenas de "
       "sistemas legados expõem APIs proprietárias que precisam ser traduzidas para o formato "
       "FHIR R4 exigido pela RNDS. A conversão manual é custosa e propensa a erros; o "
       "FHIR-Forge a automatiza com qualidade auditável e rastreabilidade completa via Langfuse.")
    _page_break(doc)


def _cap1_introducao(doc: Document) -> None:
    _h(doc, "Capítulo 1 — Introdução e Contexto")

    _h(doc, "1.1 Motivação", level=2)
    _p(doc,
       "O Brasil possui uma das redes de saúde mais complexas do mundo, com milhares de "
       "sistemas hospitalares, clínicas, laboratórios e operadoras de planos de saúde. "
       "A RNDS (Rede Nacional de Dados em Saúde), mantida pelo Ministério da Saúde via "
       "DATASUS, exige que todos os sistemas que integram com a plataforma federal utilizem "
       "o padrão FHIR R4 com o perfil brasileiro BR Core.")
    _p(doc,
       "A maioria dos sistemas legados expõe APIs REST proprietárias (OpenAPI/Swagger) sem "
       "qualquer semântica clínica padronizada. A conversão manual desses endpoints em recursos "
       "FHIR válidos demanda conhecimento especializado em terminologia clínica (SNOMED CT, "
       "LOINC, CID-10) e representa um gargalo crítico na jornada de interoperabilidade.")

    _h(doc, "1.2 Objetivos do Sistema", level=2)
    _bullet(doc, [
        "Converter automaticamente qualquer especificação OpenAPI 2.0/3.x em Bundle FHIR R4 conforme BR Core",
        "Garantir que dados PHI nunca trafeguem para LLMs cloud sem autorização explícita (LGPD)",
        "Suportar submissão direta à RNDS via autenticação mTLS",
        "Prover rastreabilidade completa de cada conversão via Langfuse e OpenTelemetry",
        "Manter cobertura de testes ≥ 80% e lint/mypy sem erros em todas as fases",
    ])

    _h(doc, "1.3 Escopo", level=2)
    _p(doc, "Dentro do escopo:", bold=True)
    _bullet(doc, [
        "Parsing de specs OpenAPI 2.0 e 3.x (YAML e JSON)",
        "Mapeamento LLM-assistido de endpoints para recursos FHIR R4",
        "Validação via HAPI FHIR $validate contra perfis BR Core",
        "Lookup terminológico (SNOMED CT, LOINC, CID-10, TUSS, SIGTAP)",
        "API REST para conversão síncrona e assíncrona",
        "Workers Dramatiq para processamento em fila",
        "Servidor MCP com 6 ferramentas FHIR para integração com Claude",
    ])
    _p(doc, "Fora do escopo:", bold=True)
    _bullet(doc, [
        "Substituição de sistemas EHR ou PEP",
        "Validação clínica semântica por profissional de saúde",
        "Interface gráfica de usuário (GUI)",
        "Certificação como dispositivo médico",
    ])

    _h(doc, "1.4 Público-alvo deste Documento", level=2)
    _p(doc,
       "Este documento é dirigido a arquitetos de solução e líderes técnicos responsáveis "
       "por avaliar, implantar ou integrar o FHIR-Forge em ecossistemas de saúde digital. "
       "O leitor deve ter familiaridade com APIs REST, containers Docker e conceitos de "
       "arquitetura de software, mas não necessariamente com FHIR ou terminologia clínica — "
       "os conceitos essenciais são explicados ao longo do documento.")

    _h(doc, "1.5 Glossário de Termos", level=2)
    glossario = [
        ("FHIR", "Fast Healthcare Interoperability Resources — padrão HL7 para troca de dados de saúde em JSON/XML/RDF"),
        ("Bundle", "Recurso FHIR que agrupa outros recursos (Patient, Observation, etc.) em um único documento"),
        ("BR Core", "Perfil brasileiro do FHIR R4, mantido pelo DATASUS/Ministério da Saúde"),
        ("RNDS", "Rede Nacional de Dados em Saúde — plataforma federal de interoperabilidade em saúde"),
        ("OpenAPI / Swagger", "Especificação de APIs REST em YAML/JSON (versões 2.0 e 3.x)"),
        ("LLM", "Large Language Model — modelo de linguagem de grande escala (ex: Claude, Llama)"),
        ("LangGraph", "Framework de agentes LLM com grafo de estados (StateGraph) da LangChain"),
        ("PHI", "Protected Health Information — dado pessoal de saúde protegido pela LGPD"),
        ("LGPD", "Lei Geral de Proteção de Dados (Lei 13.709/2018)"),
        ("mTLS", "Mutual TLS — autenticação mútua com certificado tanto no cliente quanto no servidor"),
        ("MCP", "Model Context Protocol — protocolo para LLMs chamarem ferramentas externas em tempo de execução"),
        ("SNOMED CT", "Systematized Nomenclature of Medicine — terminologia clínica internacional (licenciada pelo MS)"),
        ("LOINC", "Logical Observation Identifiers Names and Codes — terminologia de observações e exames"),
        ("CID-10", "Classificação Internacional de Doenças, 10ª revisão (sistema nacional, DATASUS)"),
        ("TUSS", "Terminologia Unificada da Saúde Suplementar — procedimentos da ANS"),
        ("SIGTAP", "Sistema de Gerenciamento da Tabela de Procedimentos do SUS (DATASUS)"),
        ("OperationOutcome", "Recurso FHIR de resultado de operação — inclui erros de validação com severity"),
        ("uv", "Gerenciador de pacotes Python da Astral — substitui pip + pip-tools + virtualenv"),
        ("Dramatiq", "Framework de filas de tarefas assíncronas para Python, com broker Redis"),
        ("Circuit Breaker", "Padrão de resiliência que abre o circuito após N falhas, evitando cascata"),
        ("SecretStr", "Tipo pydantic para strings sensíveis que nunca aparecem em repr() ou logs"),
        ("ULID", "Universally Unique Lexicographically Sortable Identifier — gerador de IDs cronológicos"),
    ]
    _table(doc, ["Termo", "Definição"], glossario, col_widths=[2.0, 5.5])
    _page_break(doc)


def _cap2_arquitetura(doc: Document) -> None:
    _h(doc, "Capítulo 2 — Arquitetura Geral")

    _h(doc, "2.1 Visão de Componentes", level=2)
    _p(doc,
       "O FHIR-Forge é organizado como um monorepo uv workspace com 7 packages Python e "
       "2 aplicações. Cada componente tem responsabilidade única e fronteiras bem definidas:")
    _table(doc,
           ["Componente", "Tipo", "Responsabilidade"],
           [
               ["packages/core",          "Package base", "Settings (pydantic), LLMRouter, tipos, exceções, logging (structlog)"],
               ["packages/swagger_lens",  "Package",      "Parser OpenAPI 2.0/3.x → SwaggerSpec tipado; flattener de $refs"],
               ["packages/fhir_forge",    "Package",      "Agente LangGraph 5 nós: parse → map → validate → fix → bundle"],
               ["packages/term_mapper",   "Package",      "Lookup SNOMED/LOINC/CID-10/TUSS com cache Redis e fallback HAPI"],
               ["packages/connectors",    "Package",      "HAPI client (retry tenacity) + RNDS client (mTLS OAuth2)"],
               ["packages/eval",          "Package",      "EvalRunner: precision/recall/semantic_diff em golden files"],
               ["packages/mcp_server",    "Package",      "6 FastMCP tools FHIR: validate, get, search, expand, concept, translate"],
               ["apps/api",               "Aplicação",    "FastAPI gateway: 4 routers + OTEL + request_id middleware"],
               ["apps/worker",            "Aplicação",    "Dramatiq actors com DLQ + idempotência Redis por job_id"],
           ],
           col_widths=[2.2, 1.5, 4.8])

    _h(doc, "2.2 Princípios Arquiteturais", level=2)
    princípios = [
        ("FHIR-first",
         "Todo dado sai validado. POST /$validate no HAPI ou validator_cli.jar local. "
         "Nenhum recurso é retornado sem validação explícita."),
        ("Async-first",
         "Toda função com I/O externo (HTTP, banco, fila) usa async def com httpx.AsyncClient. "
         "Workers Dramatiq usam asyncio.run() para orquestrar código assíncrono."),
        ("PHI-safety",
         "FEATURE_ALLOW_PHI_EGRESS=false por padrão. LLMRouter verifica o flag antes de "
         "qualquer chamada a LLM cloud. PHI é roteado obrigatoriamente para vLLM on-prem."),
        ("Observability-by-default",
         "Todo LLM call rastreado no Langfuse. Todo endpoint FastAPI tem span OTEL automático. "
         "Todos os logs via structlog com request_id, nunca print() direto."),
        ("Fail-fast, recover-gracefully",
         "tenacity retry (3x, exponential backoff) em chamadas externas. pybreaker circuit "
         "breaker em connectors. Dramatiq DLQ após 3 falhas — nenhuma tarefa perdida."),
    ]
    for titulo, desc in princípios:
        _p(doc, f"► {titulo}", bold=True, color=C_BLUE)
        _p(doc, desc)
        doc.add_paragraph()

    _h(doc, "2.3 Fluxo de Dados Principal", level=2)
    _p(doc,
       "O fluxo de uma conversão completa segue estas etapas sequenciais, executadas de forma "
       "stateful com checkpoints Postgres (via LangGraph AsyncPostgresSaver):")
    etapas = [
        ["1. Recepção", "POST /convert (API) ou fila Dramatiq", "SwaggerSpec bytes (YAML/JSON)"],
        ["2. SwaggerLens", "parse_spec() — validação + extração de endpoints", "SwaggerSpec tipado"],
        ["3. LangGraph init", "Criação do ConversionState com job_id (ULID)", "Estado inicial persistido"],
        ["4. parse_node", "Extrai operation_ids dos endpoints prioritários", "endpoints_to_process[]"],
        ["5. mapping_node", "LLM few-shot BR Core por endpoint", "fhir_resources[] (dicts JSON)"],
        ["6. validation_node", "POST /$validate HAPI por resource", "validation_errors[] ou vazio"],
        ["7. fix_node (opcional)", "LLM corrige erros (até retry_count=3)", "fhir_resources[] corrigidos"],
        ["8. bundle_node", "Monta Bundle R4 com meta.tag e warnings", "_bundle (dict final)"],
        ["9. Entrega", "Resposta JobResponse com bundle ou submissão RNDS", "Bundle FHIR R4"],
    ]
    _table(doc, ["Etapa", "Ação", "Saída"], etapas, col_widths=[1.8, 3.2, 3.5])

    _h(doc, "2.4 Tecnologias-chave", level=2)
    techs = [
        ["Python 3.12+", "Runtime principal — type hints, match statements, asyncio nativo"],
        ["LangGraph 0.6+", "Agente de conversão — StateGraph com 5 nós e roteamento condicional"],
        ["Anthropic SDK 0.50+", "Cliente LLM principal — Claude Opus 4.7 e Haiku 4.5"],
        ["FastAPI 0.115+", "Gateway REST — routers, middleware, OTEL instrumentation"],
        ["Dramatiq + Redis", "Fila de tarefas assíncronas — broker Redis, DLQ, max_retries=3"],
        ["HAPI FHIR 8.4.0", "Servidor FHIR — $validate, $expand, CRUD de recursos"],
        ["Snowstorm 7.5.0", "Servidor SNOMED CT — find_concept, ECL queries, $translate"],
        ["pydantic-settings", "Gerenciamento de configuração — SecretStr, .env, extra=ignore"],
        ["structlog + OTEL", "Observabilidade — logs estruturados JSON, spans HTTP automáticos"],
        ["tenacity + pybreaker", "Resiliência — retry exponential, circuit breaker configurable"],
        ["python-pptx / python-docx", "Geração de documentação de arquitetura (este documento)"],
    ]
    _table(doc, ["Tecnologia", "Papel no Sistema"], techs, col_widths=[2.5, 6.0])
    _page_break(doc)


def _cap3_swagger_lens(doc: Document) -> None:
    _h(doc, "Capítulo 3 — SwaggerLens: Interpretação de APIs")

    _h(doc, "3.1 Problema", level=2)
    _p(doc,
       "O mercado brasileiro de saúde utiliza specs OpenAPI em versões 2.0 (Swagger) e 3.0/3.1, "
       "frequentemente com $refs circulares, allOf/oneOf complexos e schemas sem documentação. "
       "O SwaggerLens resolve esses desafios antes de qualquer processamento LLM.")

    _h(doc, "3.2 Módulos", level=2)
    _table(doc,
           ["Módulo", "Função"],
           [
               ["parser.py",    "Valida o spec (openapi-spec-validator) e constrói SwaggerSpec tipado com endpoints"],
               ["flattener.py", "Resolve $ref recursivamente (rastreia refs visitadas para evitar loops), expande allOf/oneOf/anyOf"],
               ["extractor.py", "Infere tipos FHIR prováveis a partir de nomes de campos, tags, operationId (heurística)"],
               ["models.py",    "SwaggerSpec, Endpoint, Parameter: dataclasses com tipagem estrita"],
           ],
           col_widths=[2.0, 6.5])

    _h(doc, "3.3 Modelo de Dados", level=2)
    _table(doc,
           ["Classe", "Campos Principais"],
           [
               ["SwaggerSpec", "title: str, version: str, base_url: str, endpoints: list[Endpoint]"],
               ["Endpoint",    "path: str, method: str, operation_id: str, summary: str, params: list, request_body: dict, responses: dict"],
               ["Parameter",   "name: str, in_: str (query|path|body), required: bool, schema: dict"],
           ],
           col_widths=[2.2, 6.3])

    _h(doc, "3.4 Estratégia de Validação", level=2)
    _p(doc,
       "O parser aplica openapi-spec-validator antes de qualquer parsing. Specs inválidas "
       "levantam FhirForgeError imediatamente. Testes de propriedade via hypothesis garantem "
       "que specs válidas mínimas nunca causam exceção não tratada (200+ exemplos aleatórios).")
    _page_break(doc)


def _cap4_fhir_forge(doc: Document) -> None:
    _h(doc, "Capítulo 4 — FhirForge: O Agente LangGraph")

    _h(doc, "4.1 Visão Geral do LangGraph", level=2)
    _p(doc,
       "LangGraph é um framework para construir agentes LLM como grafos de estados dirigidos "
       "(StateGraph). Cada nó é uma função Python que recebe o estado atual e retorna um estado "
       "modificado. Arestas condicionais permitem roteamento dinâmico baseado no estado. "
       "O FHIR-Forge usa um grafo de 5 nós com loop de correção controlado por retry_count.")

    _h(doc, "4.2 ConversionState", level=2)
    _p(doc, "O estado é um TypedDict com partial=True (campos podem ser ausentes na inicialização):")
    _table(doc,
           ["Campo", "Tipo", "Descrição"],
           [
               ["swagger_spec",         "SwaggerSpec",     "Spec OpenAPI parseado pelo SwaggerLens"],
               ["endpoints_to_process", "list[str]",       "operation_ids a mapear (saída do parse_node)"],
               ["fhir_resources",       "list[dict]",      "Recursos FHIR gerados (saída do mapping_node)"],
               ["invalid_resources",    "list[dict]",      "Recursos que falharam na validação"],
               ["validation_errors",    "list[str]",       "Mensagens de erro do HAPI OperationOutcome"],
               ["retry_count",          "int",             "Contador de tentativas de fix (0–3)"],
               ["job_id",               "str",             "ULID do job — thread_id do LangGraph"],
               ["trace_id",             "str",             "ID do trace Langfuse para correlação"],
               ["warnings",             "list[str]",       "Avisos para recursos com retry_count ≥ 3"],
               ["_bundle",              "dict",            "Bundle FHIR R4 final (saída do bundle_node)"],
           ],
           col_widths=[2.2, 1.8, 4.5])

    _h(doc, "4.3 Descrição dos 5 Nós", level=2)
    nos = [
        ("parse_node",
         "Recebe o SwaggerSpec e extrai a lista de operation_ids a processar, "
         "ordenados por relevância (endpoints com request_body têm prioridade)."),
        ("mapping_node",
         "Para cada endpoint, constrói um prompt few-shot com exemplos BR Core e envia ao LLM "
         "(Claude Opus 4.7 ou vLLM on-prem se PHI). O output JSON é validado estruturalmente "
         "antes de ser adicionado a fhir_resources[]."),
        ("validation_node",
         "Faz POST /$validate no HAPI FHIR para cada recurso gerado. Recursos inválidos "
         "são movidos para invalid_resources[] e seus erros acumulados em validation_errors[]."),
        ("fix_node",
         "Envia cada recurso inválido de volta ao LLM junto com os erros de validação como "
         "contexto. O LLM gera uma versão corrigida. retry_count é incrementado. "
         "Executado apenas se retry_count < 3."),
        ("bundle_node",
         "Monta o Bundle FHIR R4 final com meta.tag fhir-forge, timestamp, e todos os recursos "
         "válidos. Recursos com retry_count ≥ 3 que ainda falham são incluídos com meta.tag "
         "de warning para inspeção manual."),
    ]
    for nome, desc in nos:
        _p(doc, nome, bold=True, color=C_BLUE)
        _p(doc, desc)
        doc.add_paragraph()

    _h(doc, "4.4 Roteamento Condicional", level=2)
    _p(doc,
       "Após validation_node, a função _route_validation() determina o próximo nó:")
    _bullet(doc, [
        "validation_errors vazio → bundle_node (sucesso)",
        "validation_errors não vazio E retry_count < 3 → fix_node (tentativa de correção)",
        "validation_errors não vazio E retry_count ≥ 3 → bundle_node (com warnings)",
    ])

    _h(doc, "4.5 Persistência de Estado", level=2)
    _p(doc,
       "O grafo é compilado com AsyncPostgresSaver (banco langgraph). Cada invocação usa "
       "thread_id=job_id, permitindo retomada automática em caso de falha de worker. "
       "O estado completo é serializado em JSON no Postgres após cada nó.")

    _h(doc, "4.6 Rastreabilidade", level=2)
    _p(doc,
       "Cada conversão gera um trace no Langfuse com input (SwaggerSpec resumido) e output "
       "(Bundle ou erro). Cada nó LLM registra tokens consumidos, latência e o prompt exato. "
       "O trace_id do Langfuse é propagado via structlog para correlação com logs de aplicação.")
    _page_break(doc)


def _cap5_llm_lgpd(doc: Document) -> None:
    _h(doc, "Capítulo 5 — Roteamento de LLMs e Política de Privacidade")

    _h(doc, "5.1 Provedores Disponíveis", level=2)
    _table(doc,
           ["Provedor", "Modelo", "Uso", "Condição de Ativação"],
           [
               ["Anthropic",   "claude-opus-4-7",               "Conversão principal (não-PHI)",    "LLM_DEFAULT_ROUTE=anthropic (padrão)"],
               ["Anthropic",   "claude-haiku-4-5-20251001",     "Tarefas rápidas / retry",           "fast=True no LLMRouter"],
               ["AWS Bedrock", "claude-sonnet-4-6-v1:0",        "Fallback cloud",                   "LLM_DEFAULT_ROUTE=bedrock"],
               ["Maritaca",    "sabia-4",                       "Textos PT-BR nativos",             "LLM_DEFAULT_ROUTE=maritaca"],
               ["vLLM on-prem","Llama 3.3 70B Instruct",        "PHI obrigatório (LGPD)",           "phi_data=True + egress=false"],
           ],
           col_widths=[1.8, 2.8, 2.2, 2.7])

    _h(doc, "5.2 Regra PHI e LGPD", level=2)
    _p(doc,
       "A variável FEATURE_ALLOW_PHI_EGRESS=false (padrão) ativa o modo PHI-safe. Quando "
       "o código chama LLMRouter.selected_provider(phi_data=True), o router retorna "
       "LLMProvider.VLLM independente de qualquer outra configuração. O dado nunca sai "
       "da infraestrutura local. Esta regra implementa o princípio da minimização de dados "
       "da LGPD (Art. 6, inciso III) e a restrição de tratamento de dados sensíveis (Art. 11).")

    _h(doc, "5.3 LLMRouter", level=2)
    _p(doc, "Classe em packages/core/src/core/llm_router.py com três métodos principais:")
    _bullet(doc, [
        "selected_provider(phi_data: bool) → LLMProvider: aplica a regra PHI e retorna o provedor",
        "model_id(provider, fast) → str: retorna o model ID configurado para o provedor",
        "get_anthropic_client(fast) → AsyncAnthropic: cria cliente com chave do SecretStr",
    ])

    _h(doc, "5.4 Segurança de Credenciais", level=2)
    _p(doc,
       "Todas as chaves de API são armazenadas como pydantic SecretStr. O método "
       "get_secret_value() é chamado apenas no momento de uso. SecretStr nunca aparece "
       "em repr(), __str__(), logs structlog ou traces Langfuse — apenas '**********'.")
    _page_break(doc)


def _cap6_term_mapper(doc: Document) -> None:
    _h(doc, "Capítulo 6 — TermMapper: Terminologia Clínica")

    _h(doc, "6.1 Sistemas de Terminologia Suportados", level=2)
    _table(doc,
           ["Sistema", "Versão / Edição", "Uso Principal"],
           [
               ["SNOMED CT", "Edição Brasil (licença MS)",   "Diagnósticos, procedimentos clínicos"],
               ["LOINC",     "Internacional (livre)",        "Observações, exames laboratoriais"],
               ["CID-10",    "OMS / DATASUS",                "Condições clínicas (sistema nacional)"],
               ["TUSS",      "ANS",                          "Procedimentos da saúde suplementar"],
               ["SIGTAP",    "DATASUS",                      "Procedimentos do SUS"],
           ],
           col_widths=[1.8, 2.5, 4.2])

    _h(doc, "6.2 Arquitetura de Lookup em Camadas", level=2)
    _p(doc,
       "O TermMapper orquestra três camadas em sequência (do mais rápido ao mais lento):")
    camadas = [
        ("1. Redis Cache", "TTL 3600s, chave '{system}:{code}', serialização JSON. 0 chamadas HTTP em cache hit."),
        ("2. Snowstorm 7.5.0", "find_concept() com ECL, $translate entre sistemas. Instância local."),
        ("3. HAPI FHIR (fallback)", "$expand para ValueSets, $lookup para codes individuais."),
    ]
    for titulo, desc in camadas:
        _p(doc, titulo, bold=True, color=C_BLUE)
        _p(doc, desc)
        doc.add_paragraph()

    _h(doc, "6.3 TermMapper.map_code()", level=2)
    _p(doc,
       "Interface principal: map_code(system, code, preferred_system) → Coding | None. "
       "Tenta cache → Snowstorm → HAPI. Persiste resultado em cache ao encontrar. "
       "Retorna None se nenhuma camada encontrar mapeamento.")
    _page_break(doc)


def _cap7_connectors(doc: Document) -> None:
    _h(doc, "Capítulo 7 — Connectors: Resiliência e Integrações Externas")

    _h(doc, "7.1 HapiFhirClient", level=2)
    _p(doc, "Operações suportadas:")
    _bullet(doc, [
        "create(resource) → dict: POST /fhir/{type}",
        "read(resource_type, resource_id) → dict: GET /fhir/{type}/{id}",
        "update(resource) → dict: PUT /fhir/{type}/{id}",
        "delete(resource_type, resource_id): DELETE /fhir/{type}/{id}",
        "search(resource_type, params) → list[dict]: GET com paginação automática (segue links 'next')",
        "validate(resource, profile) → ValidationResult: POST $validate com análise de OperationOutcome",
    ])

    _h(doc, "7.2 Retry Policy (tenacity)", level=2)
    _p(doc,
       "3 tentativas máximas, exponential backoff de 1–10 segundos, retry apenas em "
       "TransportError (falhas de rede). Erros 4xx (client errors) não são retentados — "
       "são propagados imediatamente como exceção.")

    _h(doc, "7.3 Circuit Breaker (pybreaker)", level=2)
    _p(doc,
       "Threshold de 5 falhas em 60 segundos abre o circuito. Em estado aberto, todas as "
       "chamadas falham imediatamente com CircuitBreakerError. Cooldown automático após "
       "o período configurado — o circuito testa e fecha se a chamada de teste tiver sucesso.")

    _h(doc, "7.4 RndsClient (mTLS)", level=2)
    _p(doc,
       "Autenticação com a RNDS usa certificado digital PKCS#12 (.pfx) carregado via "
       "SSLContext. O token OAuth2 é obtido do endpoint de autenticação da RNDS e "
       "cacheado com buffer de 30 segundos antes da expiração. Operações:")
    _bullet(doc, [
        "get_token() → str: obtém/renova token Bearer",
        "submit_document(document) → str: POST na RNDS, retorna número de protocolo",
        "get_document(protocol_id) → dict: consulta status de documento submetido",
    ])

    _h(doc, "7.5 Configuração RNDS", level=2)
    _table(doc,
           ["Variável", "Descrição"],
           [
               ["RNDS_BASE_URL",     "Endpoint da RNDS (homologação ou produção)"],
               ["RNDS_AUTH_URL",     "Endpoint de autenticação OAuth2"],
               ["RNDS_CERT_PATH",    "Caminho para o arquivo .pfx (em secrets/)"],
               ["RNDS_CERT_PASSWORD","Senha do certificado (SecretStr)"],
               ["RNDS_CNES",         "Código CNES do estabelecimento"],
               ["RNDS_CNPJ",         "CNPJ da organização"],
           ],
           col_widths=[2.5, 6.0])
    _page_break(doc)


def _cap8_infra(doc: Document) -> None:
    _h(doc, "Capítulo 8 — Infraestrutura: Stack Docker")

    _h(doc, "8.1 Tabela de Serviços", level=2)
    _table(doc,
           ["Serviço", "Imagem:Versão", "Porta", "Propósito"],
           [
               ["HAPI FHIR",     "hapiproject/hapi:v8.4.0-2",    "8090", "Validação FHIR, $validate, $expand, CRUD de recursos"],
               ["Snowstorm",     "snowstorm:7.5.0",              "8080", "Servidor SNOMED CT (find_concept, $translate, ECL)"],
               ["Elasticsearch", "elasticsearch:7.17.24",        "9200", "Backend de índice do Snowstorm — ES 7.x (Snowstorm 7.x usa RestHighLevelClient ES 7)"],
               ["PostgreSQL 16", "pgvector/pgvector:pg16",       "5432", "HAPI JPA, LangGraph checkpoints, dados da app"],
               ["Redis 7.4",     "redis:7.4-alpine",             "6379", "Broker Dramatiq, cache de terminologia"],
               ["MinIO",         "minio/minio:2025-09",          "9000/9001", "Armazenamento S3-compatible de artefatos"],
               ["Langfuse 2.x",  "langfuse/langfuse:2",          "3000", "Observabilidade de LLMs — Langfuse 2.x (v3 requer ClickHouse)"],
               ["ProxyLLM",      "(local build)",                "9099", "Proxy /v1/messages para dev sem API key — roteia para Ollama ou mock FHIR"],
               ["FastAPI App",   "(local build)",                "8000", "Gateway REST — /convert, /health, /fhir, /mcp"],
               ["Dramatiq",      "(local)",                      "—",    "Workers assíncronos de conversão"],
           ],
           col_widths=[1.8, 2.5, 1.0, 4.2])

    _h(doc, "8.2 Volumes e Persistência", level=2)
    _table(doc,
           ["Volume Docker", "Propósito"],
           [
               ["postgres-data", "Dados HAPI JPA, LangGraph checkpoints, tabelas da app"],
               ["es-data",       "Índice Elasticsearch para Snowstorm SNOMED"],
               ["redis-data",    "Persistência RDB/AOF do Redis (opcional)"],
               ["minio-data",    "Buckets: artifacts (bundles gerados), codesystems (dumps)"],
           ],
           col_widths=[2.0, 6.5])

    _h(doc, "8.3 Healthchecks", level=2)
    _p(doc,
       "Todos os serviços têm healthchecks configurados com start_period de 30s, "
       "interval de 10s e 12 retries. Isso garante que serviços dependentes (ex: Snowstorm "
       "que depende de Elasticsearch) só recebem tráfego após estarem healthy.")

    _h(doc, "8.4 Requisitos de Hardware", level=2)
    _table(doc,
           ["Componente", "RAM Estimada"],
           [
               ["Snowstorm + Elasticsearch", "~6 GB"],
               ["HAPI FHIR + Postgres",      "~2 GB"],
               ["Redis + MinIO",             "~512 MB"],
               ["Langfuse + App",            "~1 GB"],
               ["Total recomendado",         "12 GB livres"],
           ],
           col_widths=[3.0, 2.0])
    _p(doc,
       "Linux requer vm.max_map_count ≥ 262144 para Elasticsearch. No Docker Desktop "
       "(macOS/Windows) isso é gerenciado automaticamente.")
    _page_break(doc)


def _cap9_api(doc: Document) -> None:
    _h(doc, "Capítulo 9 — API REST: Gateway FastAPI")

    _h(doc, "9.1 Estrutura da Aplicação", level=2)
    _p(doc,
       "A aplicação é criada via create_app() com lifespan async context manager para "
       "inicialização (configure_logging, OTEL instrumentation, montagem do MCP). "
       "Middleware de request_id gera UUID para cada request e propaga via contextvars.")

    _h(doc, "9.2 Endpoints", level=2)
    _table(doc,
           ["Método", "Endpoint", "Autenticação", "Descrição"],
           [
               ["GET",  "/health",                  "Nenhuma",     "Status de HAPI, Redis e Postgres"],
               ["POST", "/convert",                 "Bearer token","Conversão síncrona ou assíncrona (async_mode)"],
               ["GET",  "/fhir/{resource_type}/{id}","Bearer token","Proxy de leitura para HAPI FHIR"],
               ["POST", "/fhir/{resource_type}/$validate","Bearer token","Proxy de validação para HAPI $validate"],
               ["*",    "/mcp",                     "Nenhuma",     "FastMCP (se FEATURE_MCP_ENABLED=true)"],
           ],
           col_widths=[1.0, 2.5, 1.8, 4.2])

    _h(doc, "9.3 Schemas de Request/Response", level=2)
    _table(doc,
           ["Schema", "Campos Principais"],
           [
               ["ConvertRequest",  "spec: str (YAML/JSON), options: ConvertOptions"],
               ["ConvertOptions",  "async_mode: bool, model: str | None, hapi_base_url: str | None"],
               ["JobResponse",     "job_id: str, status: str, bundle: dict | None, error: str | None"],
           ],
           col_widths=[2.2, 6.3])

    _h(doc, "9.4 Autenticação", level=2)
    _p(doc,
       "Bearer token validado em middleware. Com FEATURE_ADMIN_NOAUTH=true (padrão em dev), "
       "todos os endpoints admin são acessíveis sem token. Em produção, "
       "FEATURE_ADMIN_NOAUTH deve ser false.")
    _page_break(doc)


def _cap10_worker(doc: Document) -> None:
    _h(doc, "Capítulo 10 — Workers Assíncronos: Dramatiq")

    _h(doc, "10.1 Broker Redis", level=2)
    _p(doc,
       "RedisBroker configurado com a URL redis://localhost:6379. DlqMiddleware registrado "
       "no broker para monitorar mensagens exauridas. Broker configurado em configure_broker() "
       "em apps/worker/src/worker/main.py.")

    _h(doc, "10.2 convert_spec_actor", level=2)
    _p(doc,
       "Actor principal na fila 'fhir'. Parâmetros: queue_name='fhir', max_retries=3, "
       "time_limit=300_000ms (5 minutos). Recebe job_id e swagger_yaml como strings.")

    _h(doc, "10.3 Idempotência por Redis SET NX", level=2)
    _p(doc,
       "Antes de processar, o actor tenta adquirir um lock via SET NX no Redis com chave "
       "worker:lock:{job_id}. Se o lock já existe, a mensagem é descartada silenciosamente "
       "(outra instância está processando ou já processou). O lock é liberado em caso de "
       "falha para permitir retry da fila. Em sucesso, o lock persiste por 24 horas "
       "como prova de processamento.")

    _h(doc, "10.4 Dead Letter Queue (DLQ)", level=2)
    _p(doc,
       "DlqMiddleware herda de dramatiq.Middleware e monitora after_process_message. "
       "Quando retries ≥ max_retries - 1, envia automaticamente para handle_dead_letter "
       "na fila 'fhir_dlq'. O DLQ handler loga o erro, registra status FAILED no Redis "
       "e (futuramente) notifica via Langfuse.")

    _h(doc, "10.5 Fluxo Completo Assíncrono", level=2)
    fluxo = [
        ["Cliente", "POST /convert (async_mode=true)"],
        ["API",     "Enfileira convert_spec_actor.send(job_id, yaml) na fila 'fhir'"],
        ["API",     "Retorna JobResponse(job_id, status='queued', bundle=None)"],
        ["Worker",  "Consome mensagem, adquire lock NX, executa conversão"],
        ["Worker",  "Persiste resultado em Redis: worker:result:{job_id}"],
        ["Cliente", "Consulta GET /jobs/{job_id} para obter status e bundle"],
    ]
    _table(doc, ["Ator", "Ação"], fluxo, col_widths=[1.5, 7.0])
    _page_break(doc)


def _cap11_mcp(doc: Document) -> None:
    _h(doc, "Capítulo 11 — Servidor MCP: Integração com Claude")

    _h(doc, "11.1 O que é MCP (Model Context Protocol)", level=2)
    _p(doc,
       "MCP é um protocolo aberto (Anthropic) que permite que modelos LLM como Claude "
       "chamem ferramentas externas em tempo de execução de forma padronizada. O servidor "
       "MCP expõe funções Python como 'tools' que o LLM pode invocar com parâmetros "
       "estruturados, receber resultados e usar nas suas respostas.")

    _h(doc, "11.2 FastMCP", level=2)
    _p(doc,
       "FastMCP 3.x é o framework Python para servidores MCP. O FHIR-Forge usa transporte "
       "streamable-http, montado em /mcp quando FEATURE_MCP_ENABLED=true. Em modo standalone, "
       "pode ser iniciado com python -m mcp_server.server.")

    _h(doc, "11.3 Ferramentas Disponíveis", level=2)
    _table(doc,
           ["Ferramenta", "Assinatura", "Descrição", "Backend"],
           [
               ["validate_resource", "resource_json: str, profile_url: str | None → dict",
                "Valida recurso FHIR contra BR Core", "HAPI $validate"],
               ["get_resource",      "resource_type: str, resource_id: str → dict",
                "Busca recurso por tipo + ID",        "HAPI GET"],
               ["search_resources",  "resource_type: str, params: str → list[dict]",
                "Busca com paginação automática",     "HAPI GET + links next"],
               ["expand_valueset",   "valueset_url: str, filter: str | None → list[dict]",
                "Expande ValueSet FHIR",              "HAPI $expand"],
               ["find_concept",      "term: str, ecl_filter: str | None → list[dict]",
                "Busca conceito SNOMED por texto/ECL", "Snowstorm browser API"],
               ["translate_code",    "source_system, source_code, target_system → dict | None",
                "Traduz código entre sistemas",       "HAPI $translate"],
           ],
           col_widths=[2.0, 2.8, 2.2, 1.5])

    _h(doc, "11.4 Casos de Uso com Claude", level=2)
    _bullet(doc, [
        "Claude valida um recurso Patient gerado antes de enviá-lo à RNDS",
        "Claude busca o código SNOMED correto para 'diabetes mellitus tipo 2' antes de mapear",
        "Claude expande o ValueSet de raça/cor BR Core para autocomplete em formulário",
        "Claude verifica se um Organization já existe no HAPI antes de criar",
    ])
    _page_break(doc)


def _cap12_eval(doc: Document) -> None:
    _h(doc, "Capítulo 12 — Pipeline de Avaliação")

    _h(doc, "12.1 Objetivo", level=2)
    _p(doc,
       "O packages/eval mede a qualidade das conversões de forma reproduzível, comparando "
       "o output gerado contra 'golden files' (bundles FHIR esperados para cada spec de "
       "referência). Garante que nenhuma mudança de código faça regredir as conversões.")

    _h(doc, "12.2 Métricas", level=2)
    _table(doc,
           ["Métrica", "Fórmula", "Interpreta"],
           [
               ["Precision", "|tipos preditos ∩ esperados| / |tipos preditos|",
                "Dos recursos gerados, quantos estão corretos?"],
               ["Recall",    "|tipos preditos ∩ esperados| / |tipos esperados|",
                "Dos recursos esperados, quantos foram gerados?"],
               ["F1",        "2 × (P × R) / (P + R)",
                "Média harmônica entre precision e recall"],
               ["Semantic Diff", "deepdiff(pred, expected, ignore_order=True, exclude_regex_paths=...)",
                "Diferenças estruturais ignorando campos voláteis (id, meta.*)"],
           ],
           col_widths=[1.8, 3.0, 3.7])
    _p(doc,
       "Campos ignorados no diff: id, meta.lastUpdated, meta.versionId, fullUrl das entries — "
       "todos campos gerados automaticamente que variam a cada execução.")

    _h(doc, "12.3 Golden Files", level=2)
    _p(doc,
       "Localização: data/golden/swagger_specs/*.yaml e data/golden/fhir_bundles/*_expected.json. "
       "Convenção de nomenclatura: {stem}_expected.json onde stem é o nome da spec sem extensão. "
       "Para aceitar uma mudança intencional: make golden-update FILE={stem}.")

    _h(doc, "12.4 Threshold de Regressão", level=2)
    _p(doc, "Critérios de bloqueio de merge:", bold=True)
    _bullet(doc, [
        "Semantic diff não-zero em qualquer golden pair",
        "Precision < 0.80 em qualquer execução de eval",
        "HAPI $validate com errors/fatals em qualquer bundle gerado",
    ])
    _page_break(doc)


def _cap13_testes(doc: Document) -> None:
    _h(doc, "Capítulo 13 — Estratégia de Testes")

    _h(doc, "13.1 Cinco Camadas de Teste", level=2)
    _table(doc,
           ["Camada", "Marker pytest", "Velocidade", "Ferramentas", "Cob. mín."],
           [
               ["Unit",        "@pytest.mark.unit",        "< 1s cada",   "pytest, respx, unittest.mock, hypothesis", "≥ 80%"],
               ["Integration", "@pytest.mark.integration", "5–30s",       "testcontainers (Postgres + Redis + HAPI)", "≥ 80%"],
               ["Regression",  "@pytest.mark.regression",  "5–15s",       "deepdiff, golden files em data/golden/",  "—"],
               ["E2E",         "@pytest.mark.e2e",          "30–120s",     "Anthropic API real + HAPI real (6 testes validados)",      "—"],
               ["RNDS",        "@pytest.mark.rnds",         "variável",    "Sandbox RNDS (requer credenciais reais)",  "—"],
           ],
           col_widths=[1.6, 2.4, 1.5, 2.8, 1.2])

    _h(doc, "13.2 Regras para Testes Unitários", level=2)
    _bullet(doc, [
        "Sem I/O real: nenhuma chamada HTTP, banco ou filesystem",
        "Mock HTTP com respx.mock (context manager, não patch global)",
        "Mock LLM: fixture com resposta canned em JSON — nunca chamada real",
        "Property-based: hypothesis para parsers (200+ exemplos aleatórios)",
        "Um arquivo tests/unit/test_{módulo}.py por módulo",
    ])

    _h(doc, "13.3 Métricas Atuais (tag v2.1, Abril 2026)", level=2)
    _table(doc,
           ["Métrica", "Valor"],
           [
               ["Testes unitários passando", "289"],
               ["Testes integration (HAPI testcontainer)", "3 (Bundle/$validate sem fatais nem dom-6)"],
               ["Testes regression",          "7 (golden pairs + invariantes estruturais)"],
               ["Testes e2e validados",       "6 (pipeline real Anthropic + HAPI)"],
               ["Cobertura de código",       "91%"],
               ["ruff check",                "0 erros"],
               ["mypy strict",               "0 erros"],
               ["Hipótese (hypothesis)",     "> 200 exemplos por parser"],
           ],
           col_widths=[3.0, 4.0])

    _h(doc, "13.4 CI Planejado (GitHub Actions)", level=2)
    _table(doc,
           ["Gatilho", "Testes", "Tempo alvo"],
           [
               ["push",   "unit + lint + mypy",                    "< 3 min"],
               ["PR",     "unit + integration + regression",        "< 12 min"],
               ["main",   "tudo (unit + int + regression + e2e)",   "< 20 min"],
           ],
           col_widths=[1.5, 4.5, 2.5])
    _page_break(doc)


def _cap14_seguranca(doc: Document) -> None:
    _h(doc, "Capítulo 14 — Segurança")

    _h(doc, "14.1 Superfície de Ataque", level=2)
    _table(doc,
           ["Superfície", "Risco", "Mitigação"],
           [
               ["API REST pública",  "Acesso não autorizado",     "Bearer token; FEATURE_ADMIN_NOAUTH=false em produção"],
               ["LLM ingress",       "Data exfiltration (PHI)",   "FEATURE_ALLOW_PHI_EGRESS=false; vLLM on-prem para PHI"],
               ["RNDS egress",       "Intercepção de dados",      "mTLS com certificado digital PKCS#12"],
               ["Armazenamento",     "Vazamento de secrets",      "SecretStr pydantic; .env em .gitignore; secrets/ fora do repo"],
               ["LLM prompts",       "Prompt injection",          "Saída JSON structured; validação de schema antes de usar"],
           ],
           col_widths=[2.0, 2.5, 4.0])

    _h(doc, "14.2 Gestão de Secrets", level=2)
    _bullet(doc, [
        "Todas as chaves sensíveis usam pydantic SecretStr — mascaradas em repr/logs",
        "Arquivo .env listado em .gitignore — nunca commitado",
        "Certificados RNDS (.pfx) em secrets/ — diretório excluído do repo",
        "Em produção: usar variáveis de ambiente (Kubernetes Secrets ou AWS SSM)",
    ])

    _h(doc, "14.3 LGPD e PHI", level=2)
    _p(doc,
       "Dados de saúde são 'dados pessoais sensíveis' sob a LGPD (Art. 5, inciso II). "
       "O tratamento desses dados exige base legal (Art. 11) e adoção de medidas de segurança "
       "adequadas (Art. 46). O FHIR-Forge implementa privacidade por design: PHI nunca "
       "egride para serviços cloud sem autorização explícita da organização (flag de configuração "
       "que deve ser revisado com DPO antes de habilitar).")
    _page_break(doc)


def _cap15_operacoes(doc: Document) -> None:
    _h(doc, "Capítulo 15 — Operações")

    _h(doc, "15.1 Makefile — Principais Targets", level=2)
    _table(doc,
           ["Target", "Descrição"],
           [
               ["make bootstrap",    "Setup completo de zero (idempotente)"],
               ["make up",           "Sobe toda a stack Docker"],
               ["make health",       "Verifica endpoints de todos os serviços"],
               ["make test",         "Roda 289 testes unitários"],
               ["make test-int",     "Testes de integração (testcontainers)"],
               ["make test-e2e",     "Testes e2e (stack up + ANTHROPIC_API_KEY)"],
               ["make lint",         "ruff check + mypy strict"],
               ["make fmt",          "ruff format + fix automático"],
               ["make api",          "FastAPI em modo dev (reload)"],
               ["make worker",       "Dramatiq worker"],
               ["make proxy",        "ProxyLLM local (porta 9099) — dev sem API key"],
               ["make docs",         "Gera PPTX + DOCX de arquitetura em docs/"],
               ["make validate",     "Valida FHIR local via validator_cli.jar"],
               ["make db-migrate",   "Alembic upgrade head"],
               ["make logs-hapi",    "tail -f logs do HAPI FHIR"],
               ["make nuke",         "APAGA todos os volumes Docker (⚠ irreversível)"],
           ],
           col_widths=[2.2, 6.3])

    _h(doc, "15.2 Monitoramento", level=2)
    _bullet(doc, [
        "Langfuse (http://localhost:3000): traces de LLM, tokens, latência por conversão",
        "OTEL spans: FastAPI endpoint → conversão → validação HAPI (Jaeger ou OTLP)",
        "structlog JSON: todos os logs com request_id, job_id, resource_type",
        "Dramatiq admin: via redis-cli para inspecionar filas",
    ])

    _h(doc, "15.3 Troubleshooting Rápido", level=2)
    _table(doc,
           ["Sintoma", "Causa Provável", "Solução"],
           [
               ["elasticsearch reinicia em loop", "vm.max_map_count baixo",            "sudo sysctl -w vm.max_map_count=262144"],
               ["hapi 503 em /fhir/metadata",     "Postgres ainda subindo",            "Aguardar 30s; make health"],
               ["langfuse 502",                   "NEXTAUTH_SECRET vazio",             "make secrets; atualizar .env"],
               ["uv sync falha em torch",         "Pouco espaço em /tmp",              "UV_CACHE_DIR=/path/grande uv sync"],
               ["convert retorna 422",            "Spec OpenAPI inválida",             "openapi-spec-validator spec.yaml"],
               ["DLQ crescendo",                  "HAPI indisponível ou LLM offline",  "make health; verificar chave API"],
           ],
           col_widths=[2.5, 2.5, 3.5])

    _h(doc, "15.4 Backup e Recuperação", level=2)
    _bullet(doc, [
        "Postgres: pg_dump --format=custom forge > backup.dump (via make db-shell)",
        "MinIO: mc mirror minio/artifacts ./backup-artifacts",
        "Redis: BGSAVE gera dump.rdb (montado em redis-data volume)",
        "Langfuse: exportar traces via API antes de nuke",
    ])
    _page_break(doc)


def _apendices(doc: Document) -> None:
    _h(doc, "Apêndice A — Dependências Python Principais")
    _table(doc,
           ["Pacote", "Versão mín.", "Finalidade"],
           [
               ["anthropic",          "0.50",    "Cliente LLM Claude (Anthropic e Bedrock)"],
               ["fastapi",            "0.115",   "Framework web assíncrono"],
               ["langfuse",           "2.x",     "Rastreamento de chamadas LLM (v3 requer ClickHouse)"],
               ["langgraph",          "0.6",     "Agente de conversão (StateGraph)"],
               ["langgraph-checkpoint-postgres","3.0","Persistência de estado LangGraph"],
               ["openapi-pydantic",   "0.5",     "Parsing tipado de specs OpenAPI"],
               ["openapi-spec-validator","0.7",  "Validação de specs antes de parsear"],
               ["pydantic",           "2.9",     "Validação e serialização de dados"],
               ["pydantic-settings",  "2.6",     "Gerenciamento de configuração via env"],
               ["pybreaker",          "1.2",     "Circuit breaker para chamadas externas"],
               ["redis",              "5.2",     "Client Redis (cache + broker Dramatiq)"],
               ["respx",              "0.21",    "Mock de HTTP (httpx) em testes"],
               ["structlog",          "24.4",    "Logging estruturado JSON"],
               ["tenacity",           "9.0",     "Retry com backoff para I/O"],
               ["python-pptx",        "1.0",     "Geração de apresentações PPTX"],
               ["python-docx",        "1.1",     "Geração de documentos DOCX (este documento)"],
           ],
           col_widths=[2.5, 1.5, 4.5])
    _page_break(doc)

    _h(doc, "Apêndice B — Variáveis de Ambiente Principais")
    _table(doc,
           ["Variável", "Padrão", "Descrição", "Sensível"],
           [
               ["ANTHROPIC_API_KEY",         "",                            "Chave API Anthropic",                     "Sim"],
               ["ANTHROPIC_MODEL_PRIMARY",   "claude-opus-4-7",             "Modelo LLM principal",                    "Não"],
               ["ANTHROPIC_MODEL_FAST",      "claude-haiku-4-5-20251001",   "Modelo LLM rápido",                       "Não"],
               ["LLM_DEFAULT_ROUTE",         "anthropic",                   "Provedor LLM padrão (anthropic/bedrock/maritaca/vllm)", "Não"],
               ["FEATURE_ALLOW_PHI_EGRESS",  "false",                       "Permite PHI em LLMs cloud (LGPD)",        "Não"],
               ["FEATURE_MCP_ENABLED",       "true",                        "Habilita servidor MCP em /mcp",           "Não"],
               ["FEATURE_ADMIN_NOAUTH",      "true",                        "Bypass auth admin (DEV ONLY)",            "Não"],
               ["HAPI_BASE_URL",             "http://localhost:8090/fhir",  "Endpoint do HAPI FHIR",                   "Não"],
               ["SNOWSTORM_BASE_URL",        "http://localhost:8080",       "Endpoint do Snowstorm",                   "Não"],
               ["POSTGRES_PASSWORD",         "forge_dev",                   "Senha do Postgres",                       "Sim"],
               ["RNDS_CERT_PATH",            "",                            "Caminho para certificado .pfx RNDS",      "Não"],
               ["RNDS_CERT_PASSWORD",        "",                            "Senha do certificado RNDS",               "Sim"],
               ["LANGFUSE_PUBLIC_KEY",       "",                            "Chave pública Langfuse",                  "Não"],
               ["LANGFUSE_SECRET_KEY",       "",                            "Chave secreta Langfuse",                  "Sim"],
               ["LOG_LEVEL",                 "DEBUG",                       "Nível de log (DEBUG/INFO/WARNING)",       "Não"],
               ["LOG_FORMAT",                "console",                     "Formato de log (console/json)",           "Não"],
           ],
           col_widths=[2.5, 2.0, 2.8, 1.2])
    _page_break(doc)

    _h(doc, "Apêndice C — Referências")
    refs = [
        ("FHIR R4 Specification (HL7)",           "https://hl7.org/fhir/R4/"),
        ("BR Core — Perfil FHIR Brasileiro",      "https://www.saude.gov.br/fhir/r4"),
        ("RNDS — Guia de Integração",             "https://rnds-guia.saude.gov.br/"),
        ("HAPI FHIR Server",                      "https://hapifhir.io/"),
        ("Snowstorm SNOMED CT Server",            "https://github.com/IHTSDO/snowstorm"),
        ("LangGraph Framework",                   "https://langchain-ai.github.io/langgraph/"),
        ("FastMCP — Model Context Protocol",      "https://github.com/jlowin/fastmcp"),
        ("LGPD — Lei 13.709/2018",                "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13709.htm"),
        ("uv — Python Package Manager",           "https://docs.astral.sh/uv/"),
        ("Dramatiq — Task Queue",                 "https://dramatiq.io/"),
    ]
    for titulo, url in refs:
        _p(doc, f"{titulo}")
        _p(doc, url, italic=True, color=C_BLUE)
        doc.add_paragraph()


# ── Main ─────────────────────────────────────────────────────────────────────

def main() -> None:
    doc = Document()
    _configure_document(doc)

    _capa(doc)
    _cap1_introducao(doc)
    _cap2_arquitetura(doc)
    _cap3_swagger_lens(doc)
    _cap4_fhir_forge(doc)
    _cap5_llm_lgpd(doc)
    _cap6_term_mapper(doc)
    _cap7_connectors(doc)
    _cap8_infra(doc)
    _cap9_api(doc)
    _cap10_worker(doc)
    _cap11_mcp(doc)
    _cap12_eval(doc)
    _cap13_testes(doc)
    _cap14_seguranca(doc)
    _cap15_operacoes(doc)
    _apendices(doc)

    Path("docs").mkdir(exist_ok=True)
    out = Path("docs") / "fhir-forge-arquitetura.docx"
    doc.save(str(out))
    print(f"Gerado: {out}  (15 capítulos + apêndices)")


if __name__ == "__main__":
    main()
